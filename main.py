from fastapi import FastAPI,HTTPException,Request
from firestore import  get_all_documents_suratKeluar,get_all_documents,get_documents_by_title_suratKeluar,put_documents_by_title_suratKeluar,save_to_firestore,get_document_by_id,get_documents_by_title, put_documents_by_title, save_to_firestore_suratKeluar
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, RedirectResponse, FileResponse, StreamingResponse
from starlette.middleware.sessions import SessionMiddleware
from docx import Document
import io


app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://127.0.0.1:5500"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

#STATIC_ID = "welldone"
#STATIC_PASSWORD = "alex1234"


@app.get("/")
def read_root():
    return {"server Sukses"}


@app.post("/save")
async def add(request: Request):
    data = await request.json()
    
    # Extracting the data
    No = data.get("No")
    Dari = data.get("Dari")
    TanggalSurat = data.get("TanggalSurat")
    TerimaTanggal = data.get("TerimaTanggal")
    Nomorsurat = data.get("Nomorsurat")
    Perihal = data.get("Perihal")
    IsiDisposisi = data.get("IsiDisposisi")
    Keterangan = data.get("Keterangan")
    
    try:
        # Prepare data for saving
        simpan = {
            "No": No,
            "Dari": Dari,
            "Tanggal_Surat": TanggalSurat,
            "Tanggal_Terima": TerimaTanggal,
            "Nomor_Surat": Nomorsurat,
            "perihal": Perihal,
            "Isi_Disposisi": IsiDisposisi,
            "Keterangan": Keterangan
        }
        
        # Save to Firestore
        if not save_to_firestore(simpan):
            raise HTTPException(status_code=500, detail="Gagal menyimpan data ke Firestore")
        
        # Load the Word template
        template_path = "dipodisiTemp.docx"
        doc = Document(template_path)
        
        # Replace placeholders with actual data
        def replace_placeholder(doc, placeholder, value):
            for para in doc.paragraphs:
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, value)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
        
        # Replace template placeholders with the saved data
        replace_placeholder(doc, "{{ No }}", No)
        replace_placeholder(doc, "{{ Dari }}", Dari)
        replace_placeholder(doc, "{{ Tanggal_Surat }}", TanggalSurat)
        replace_placeholder(doc, "{{ Tanggal_Terima }}", TerimaTanggal)
        replace_placeholder(doc, "{{ Nomor_Surat }}", Nomorsurat)
        replace_placeholder(doc, "{{ perihal }}", Perihal)
        replace_placeholder(doc, "{{ Keterangan }}", Keterangan)
        
        # Save the modified document to a buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Return the Word document as a download response
        headers = {
            'Content-Disposition': 'attachment; filename="disposisi_output.docx"'
        }
        return StreamingResponse(buffer, headers=headers, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    except ValueError as e:
        return {"error": str(e)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal upload file: {e}")

#pending dulu dipake kalo perlu aja
#----------------------------------------------------------------
#@app.get("/fine")
#async def findDocuments(id: str):
#    dokumen = get_document_by_id(id)
#    try:
#        if dokumen:
#            return {"document": dokumen}
#        else:
#            raise HTTPException(status_code=404, detail="dokumen tidak terdaftar")
#    except Exception as e:
#        raise HTTPException(status_code=500, detail=str(e))
#----------------------------------------------------------------

@app.get("/CariArsip/")
async def findDoc_judul(NomorSurat: str):
    try:
        dokumen = get_documents_by_title(NomorSurat)
        if dokumen:
            return {"documents": dokumen}
        else:
            raise HTTPException(status_code=404, detail="Dokumen tidak ditemukan")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/edit")
async def edit_document(request: Request):
    try:
        data = await request.json()
        No = data.get("No")
        Dari = data.get("Dari")
        TanggalSurat = data.get("TanggalSurat")
        TerimaTanggal = data.get("TerimaTanggal")
        Nomorsurat = data.get("Nomorsurat")
        Perihal = data.get("Perihal")
        IsiDisposisi = data.get("IsiDisposisi")
        Keterangan = data.get("Keterangan")

        update_data = {}
        if No:
            update_data['No'] = No
        if Dari:
            update_data['Dari'] = Dari
        if TanggalSurat:
            update_data['Tanggal_Surat'] = TanggalSurat
        if TerimaTanggal:
            update_data['Tanggal_Terima'] = TerimaTanggal
        if Nomorsurat:
            update_data['Nomor_Surat'] = Nomorsurat
        if Perihal:
            update_data['perihal'] = Perihal
        if IsiDisposisi:
            update_data['Isi_Disposisi'] = IsiDisposisi
        if Keterangan:
            update_data['Keterangan'] = Keterangan

        if not update_data:
            raise HTTPException(status_code=400, detail="Tidak ada data yang diberikan untuk diperbarui")

        if put_documents_by_title(Nomorsurat, update_data):
            return {"detail": "Dokumen berhasil diperbarui"}
        else:
            raise HTTPException(status_code=404, detail="Dokumen tidak ditemukan")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/showAll")
async def showAll():
    try:
        # Mendapatkan semua dokumen dari Firestore
        documents = get_all_documents()
        if not documents:
            raise HTTPException(status_code=404, detail="dokument tidak ditemukan")
        
        # Mengembalikan hasil sebagai JSON response
        return JSONResponse(content={"documents": documents})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred: {e}")
    
#------------------------------------------------
#untuk arsip keluar
#------------------------------------------------

@app.post("/saveKeluar")
async def add(request: Request):
    data = await request.json()
    No = data.get("nomorSurat")
    kepada = data.get("kepada")
    TanggalSurat = data.get("tanggalSurat")
    perihal = data.get("perihal")
    keterangan = data.get("keterangan")

    try:    
        simpan = {
            "Nomor_surat": No,
            "Kepada": kepada,
            "Tanggal_Surat": TanggalSurat,
            "perihal": perihal,
            "Keterangan": keterangan
        }
        
        if not save_to_firestore_suratKeluar(simpan):
            raise HTTPException(status_code=500, detail="Gagal menyimpan data ke Firestore")
        
    except ValueError as e:
        return {"error": str(e)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal upload file: {e}")
    
    return {"berhasil menambahkan": No}

@app.put("/editSuratKeluar")
async def edit_document(request: Request):
    try:
        data = await request.json()
        No = data.get("nomorSurat")
        kepada = data.get("kepada")
        TanggalSurat = data.get("tanggalSurat")
        perihal = data.get("perihal")
        keterangan = data.get("keterangan")

        update_data = {}
        if No:
            update_data['Nomor_surat'] = No
        if kepada:
            update_data['Kepada'] = kepada
        if TanggalSurat:
            update_data['Tanggal_Surat'] = TanggalSurat
        if perihal:
            update_data['perihal'] = perihal
        if keterangan:
            update_data['Keterangan'] = keterangan

        if not update_data:
            raise HTTPException(status_code=400, detail="Tidak ada data yang diberikan untuk diperbarui")

        if put_documents_by_title_suratKeluar(No, update_data):
            return {"detail": "Dokumen berhasil diperbarui"}
        else:
            raise HTTPException(status_code=404, detail="Dokumen tidak ditemukan")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/CariArsipKeluar/")
async def findDoc_judul(NomorSurat: str):
    try:
        dokumen = get_documents_by_title_suratKeluar(NomorSurat)
        if dokumen:
            return {"documents": dokumen}
        else:
            raise HTTPException(status_code=404, detail="Dokumen tidak ditemukan")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
@app.get("/showAllKeluar")
async def showAll():
    try:
        # Mendapatkan semua dokumen dari Firestore
        documents = get_all_documents_suratKeluar()
        if not documents:
            raise HTTPException(status_code=404, detail="dokument tidak ditemukan")
        
        # Mengembalikan hasil sebagai JSON response
        return JSONResponse(content={"documents": documents})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred: {e}")

@app.post("/download-disposisi/")
async def download_disposisi(request: Request):
    # Ambil data dari request
    data = await request.json()
    No = data.get("No")
    Dari = data.get("Dari")
    Tanggal_Surat = data.get("Tanggal_Surat")
    Tanggal_Terima = data.get("Tanggal_Terima")
    Nomor_Surat = data.get("Nomor_Surat")
    perihal = data.get("perihal")
    Keterangan = data.get("Keterangan")
    
    # Load template Word
    template_path = "dipodisiTemp.docx"
    doc = Document(template_path)
    
    # Fungsi untuk mengganti semua placeholder di dalam template
    def replace_placeholder(doc, placeholder, value):
        for para in doc.paragraphs:
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, value)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)
    
    # Ganti nilai placeholder di dalam template
    replace_placeholder(doc, "{{ No }}", No)
    replace_placeholder(doc, "{{ Dari }}", Dari)
    replace_placeholder(doc, "{{ Tanggal_Surat }}", Tanggal_Surat)
    replace_placeholder(doc, "{{ Tanggal_Terima }}", Tanggal_Terima)
    replace_placeholder(doc, "{{ Nomor_Surat }}", Nomor_Surat)
    replace_placeholder(doc, "{{ perihal }}", perihal)
    replace_placeholder(doc, "{{ Keterangan }}", Keterangan)
    
    # Simpan dokumen yang sudah diubah ke dalam buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Kembalikan file sebagai respons download menggunakan StreamingResponse
    headers = {
        'Content-Disposition': 'attachment; filename="disposisi_output.docx"'
    }
    return StreamingResponse(buffer, headers=headers, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')